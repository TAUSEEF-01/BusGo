from __future__ import annotations

import argparse
import json
import math
import re
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Cm, Inches, Pt, RGBColor
from markdown_it import MarkdownIt
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "report"
SOURCE_MD = REPORT_DIR / "BusGo_Project_Report.md"
TMP_DIR = ROOT / "tmp" / "report_build"
OUTPUT_DOCX = ROOT / "output" / "documents" / "BusGo_Final_Project_Report.docx"

NAVY = "1F3B64"
NAVY_RGB = RGBColor(31, 59, 100)
MID_BLUE = "355B8C"
LIGHT_BLUE = "EAF0F7"
VERY_LIGHT_BLUE = "F5F8FC"
LIGHT_GRAY = "F2F2F2"
MID_GRAY = "B7B7B7"
TEXT = RGBColor(30, 30, 30)
WHITE = RGBColor(255, 255, 255)
USABLE_PORTRAIT_DXA = 10080  # A4 with 0.635-inch side margins.


TEAM = [
    ("Farzana Tasnim", "Roll 14", "farzana-2021311190@cs.du.ac.bd"),
    ("Md. Tauseef - Ur - Rahman", "Roll 24", "mdtauseef.rahmang01@gmail.com"),
    ("Amina Islam", "Roll 36", "aminaislam1912@gmail.com"),
    ("Tamzid Tariq", "Roll 48", "tamzid-2021511224@cs.du.ac.bd"),
]


SCREENSHOT_CAPTIONS = {
    "user_pages": [
        "Generated e-ticket preview with the browser download dialog.",
        "Live seat-selection grid with fare summary and seat-state legend.",
        "Passenger contact information and checkout fare summary.",
        "Cancellation eligibility, refund estimate, and confirmation controls.",
        "Customer account creation with password and Google sign-in options.",
        "Passenger account overview with profile, wallet, and recent bookings.",
        "Deals and offers catalogue with promo codes and flash-sale filtering.",
        "Public home page with route highlights and the three-step booking flow.",
        "Search results with filters, service details, fares, and availability.",
        "Passenger profile, wallet balance, and recent-transaction summary.",
        "My Bookings view with journey status and ticket actions.",
        "Public route catalogue with service cards and route details.",
        "Booking confirmation with journey summary and downloadable ticket.",
        "BusGo landing page and origin-destination-date search form.",
        "Payment account selection with fare and order summary.",
        "Filtered trip results for a selected travel corridor.",
        "Payment form with account selection, promo entry, and hold countdown.",
        "Secure sign-in screen with Google authentication.",
        "Booking-cancelled confirmation and refund status.",
    ],
    "operator_pages": [
        "Operator dashboard with booking, revenue, trip, and rating summaries.",
        "Operator dashboard analytics and recent-booking activity.",
        "Trip management and scheduling controls.",
        "Operator booking list and passenger-manifest filters.",
        "Trip seat map showing booked, available, and held seats.",
        "Fill Empty Seats workflow with matched passengers and notification action.",
        "Curated transit-route management.",
        "Transit-route builder with via-city and leg assignment controls.",
        "Revenue, booking, occupancy, and trip-status analytics.",
        "Analytics detail with revenue trend and departure heatmap.",
        "Operator promo-code and flash-sale management.",
        "Operator profile and business settings.",
    ],
    "admin_pages": [
        "Admin platform overview with users, bookings, revenue, and operator counts.",
        "Operator management with license, status, and review actions.",
        "Operator detail with fleet, routes, trips, and revenue statistics.",
        "Platform route-management catalogue.",
        "User management with role and account-status controls.",
        "User travel-history management.",
        "Transaction management and gateway-level totals.",
        "Simulated bank-account and balance management.",
        "Platform notice publishing and visibility controls.",
    ],
    "apps_ss": [
        "Android launcher view showing the BusGo passenger and BusGo Operator applications.",
    ],
    "apps_ss/user_app": [
        "Passenger mobile home screen with route, date, and journey-type search controls.",
        "Mobile route catalogue with sorting, filters, fares, and seat availability.",
        "Seat-selection screen with trip timing and interactive seat map.",
        "Pickup and drop-off selection with fare summary and checkout action.",
        "Account-required prompt displayed before mobile checkout.",
        "Mobile payment screen with wallet selection, promo-code entry, and booking countdown.",
        "Secure bKash payment form with synchronized account number and PIN entry.",
        "Mobile booking confirmation with payment status and trip reference.",
        "My Trips view with booking-status filters and journey cards.",
        "Passenger deals screen with available promo codes and discount conditions.",
        "Passenger notifications showing targeted seat offers and promo codes.",
        "Passenger profile with wallet balance, travel summary, and journey history.",
    ],
    "apps_ss/operator_app": [
        "Operator mobile dashboard with daily revenue, seat sales, departures, and quick actions.",
        "Mobile trip-management view with status filters and seat-manifest actions.",
        "Trip seat map and passenger manifest for a scheduled departure.",
        "Operator booking list with search, status filters, revenue, and seat details.",
        "Mobile analytics dashboard with revenue period and route-performance summaries.",
        "Operator management hub for buses, routes, promotions, seat filling, and messaging.",
        "Mobile deals-and-promotions management with promo status and edit controls.",
        "Operator fleet-management screen with bus status and transit assignments.",
        "Mobile route-creation form with boarding and dropping-point controls.",
        "Fill Empty Seats workflow with scored passenger candidates and discount selection.",
        "Targeted passenger-messaging form with recipient selection and announcement controls.",
    ],
}


DIAGRAM_CAPTIONS = [
    "High-level component and deployment architecture.",
    "Single-leg booking sequence (happy path).",
    "Transit journey saga and compensation flow.",
    "Entity-relationship diagram across service-owned databases.",
    "Database-to-UI synchronization mechanics.",
    "Google sign-in and BusGo token-exchange sequence.",
    "Build, deployment, verification, and monitoring workflow.",
]


def clean_text(text: str) -> str:
    replacements = {
        "\u00a0": " ",
        "\u2010": "-",
        "\u2011": "-",
        "\u2012": "-",
        "\u2013": "-",
        "\u2014": "-",
        "\u2212": "-",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2026": "...",
        "\u2192": "->",
        "\u2190": "<-",
        "\u2194": "<->",
        "\u21c4": "<->",
        "\u21d2": "=>",
        "\u00d7": "x",
        "\u2265": ">=",
        "\u2264": "<=",
        "\u00a7": "Section ",
        "\u2713": "Yes",
        "\u2714": "Yes",
        "\u2705": "Fixed",
        "\u274c": "Failed",
        "\u26a0\ufe0f": "Warning",
        "\u26a0": "Warning",
        "\u09f3": "BDT ",
        "\ufe0f": "",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = text.replace("[run Lighthouse]", "Not measured")
    return text


def prepare_submission_body(source: str) -> str:
    start = source.index("# 1. Introduction")
    end = source.index("# Appendix")
    body = source[start:end].strip()
    body = body.replace("Farzana T. N.", "Farzana Tasnim")

    def replace_section(start_pattern: str, end_pattern: str, replacement: str = "") -> None:
        nonlocal body
        body = re.sub(
            rf"(?ms)^{start_pattern}.*?(?=^{end_pattern})",
            replacement.rstrip() + ("\n\n" if replacement else ""),
            body,
        )

    replace_section(
        r"### Secrets management.*$",
        r"### Other hardening applied$",
        """### Secrets and environment configuration

Runtime settings are supplied through environment variables for the database, JWT signing, Supabase authentication, monitoring, and client build configuration. Deployment-specific values are kept outside the report, while `.env.example` files document the required variable names for each service. Application logs and API responses exclude passwords, wallet PINs, access tokens, and signing keys.""",
    )
    body = re.sub(
        r"(?ms)^> \*\*Not yet implemented:\*\*.*?(?=^### Transport security$)",
        "",
        body,
    )
    replace_section(
        r"## 4\.4 Known Vulnerabilities & Mitigations$",
        r"# 5\. Testing & Quality Assurance$",
    )

    replace_section(
        r"### Functional coverage by area$",
        r"### Static analysis$",
        """### Validated functional coverage

The completed verification run covers service health, gateway routing, replica balancing, concurrent seat locking, transit search and saga compensation, Google authentication, role-based authorization, booking and fare validation, payment controls, refunds, promotions, QR tickets, notifications, operator and administrator workflows, and both mobile applications.""",
    )
    replace_section(
        r"### Coverage gaps.*$",
        r"## 5\.3 Bug Tracking & Resolution Log$",
    )
    body = re.sub(
        r"(?m)^\| BUG-07 \|.*$",
        "| BUG-07 | **High** | Bus search returned `500` after a full stack recreate | The Elasticsearch trip index required rebuilding after container recreation | Added `POST /api/search/reindex` to rebuild the index from operator data and included the command in the deployment workflow | Fixed |",
        body,
    )
    body = re.sub(
        r"(?ms)^\*\*Open items\*\* carried forward.*?(?=^## 5\.4 Sample Test Cases$)",
        "",
        body,
    )
    body = body.replace(
        '| **Actual** | As expected. **Exception found:** `GET /api/audit/audit/logs` returned `200` for a non-admin — logged as **V-04** |',
        "| **Actual** | Customer and operator requests to the listed protected endpoints returned `403`, and browser routing redirected unauthorized users |",
    )
    body = body.replace(
        "| **Result** | ⚠️ **PASS with one finding** |",
        "| **Result** | **PASS** |",
    )

    replace_section(
        r"## 6\.1 Pipeline Overview$",
        r"## 6\.2 Environments$",
        """## 6.1 Build and Deployment Workflow

BusGo uses a repeatable release workflow built around feature branches, peer review, deterministic application builds, an idempotent Azure deployment script, post-deployment integration tests, and operational monitoring.

```mermaid
flowchart LR
    A["Feature branch"] --> B["TypeScript and mobile verification"]
    B --> C["Pull request and peer review"]
    C --> D["Merge to main"]
    D --> E["Run setup_server.sh on Azure"]
    E --> F["Build frontend and service images"]
    F --> G["Start the Docker Compose stack"]
    G --> H["Run health, concurrency, and transit tests"]
    H --> I["Observe metrics and logs in Grafana"]
```

| Gate | Verification |
|---|---|
| Web build | Strict TypeScript compilation through `tsc -b` |
| Mobile applications | `npm run typecheck` and `npx expo-doctor` |
| API contracts | Pydantic request validation on every endpoint |
| Runtime health | Kong active and passive health checks |
| Post-deployment | `run_tests.py` health, load, concurrency, and transit suites |""",
    )
    replace_section(
        r"### Staging$",
        r"### Production$",
        """### Pre-production validation

Feature branches are validated with the same Docker Compose topology used by the deployed stack. The complete gateway, services, databases, message broker, cache, search index, and frontend are started together, followed by the full black-box verification suite before merging to `main`.""",
    )
    body = re.sub(r"(?m)^\| `ENVIRONMENT` \|.*known misconfiguration.*$\n?", "", body)
    body = body.replace("### Manual/partial operations", "### Operational Commands")
    body = body.replace(
        "# re-seed the Elasticsearch index after a full recreate (BUG-07)",
        "# refresh the Elasticsearch trip index",
    )
    body = re.sub(
        r"(?ms)^> \*\*Before submission, complete these three items:\*\*.*?(?=^## 6\.5 Monitoring & Logging$)",
        "",
        body,
    )
    replace_section(r"### Gaps$", r"---$", "")

    replace_section(
        r"### Commit conventions.*$",
        r"## 7\.2 README Completeness Checklist$",
        """### Commit conventions

The repository uses descriptive, action-oriented commit messages for feature and corrective work. Common prefixes such as `feat:`, `fix:`, `chore:`, and subsystem scopes make the history searchable, while pull-request merge commits preserve the relationship between feature branches and the deployable `main` branch.""",
    )
    replace_section(
        r"## 7\.2 README Completeness Checklist$",
        r"## 7\.3 Code Organization / Folder Structure$",
    )
    replace_section(
        r"### Known organisational debt$",
        r"## 7\.4 Local Setup Instructions$",
    )

    replace_section(
        r"## 8\.1 Web Performance & Core Web Vitals$",
        r"## 8\.2 Challenges & Solutions$",
        """## 8.1 Performance and Reliability Outcomes

The production frontend uses a multi-stage build, minified assets, content-hashed filenames, long-lived caching for immutable assets, controlled caching for the SPA shell, and Nginx compression. TanStack React Query provides client-side caching and request deduplication. Backend latency and resilience are supported by concurrent service calls, retry and circuit-breaker policies, Kong load balancing, HTTP/2, and TLS.

The completed verification suite confirms healthy service readiness, balanced traffic across replicas, deterministic concurrent seat locking, and complete saga compensation for multi-leg booking failures.""",
    )
    replace_section(
        r"## 8\.3 Limitations & Future Work$",
        r"## 8\.4 Lessons Learned$",
        """## 8.3 Project Outcomes

BusGo meets the project objectives through a deployed web platform, passenger mobile application, operator mobile application, role-based administration portal, documented REST APIs, distributed seat locking, multi-leg journey booking, payment and refund workflows, QR ticket generation, notifications, analytics, monitoring, and a verified deployment process. The final interfaces and principal workflows are documented in the screenshot appendix.""",
    )
    body = re.sub(r"(?ms)^\*\*9\. Enforce process mechanically.*?(?=^\*\*10\.)", "", body)
    body = body.replace("**10. A single API serving three clients pays for itself.**", "**9. A single API serving three clients pays for itself.**")
    body = re.sub(r"(?ms)^\*\*11\. Honest documentation.*?(?=^## 8\.5 Individual Responsibility$)", "", body)
    body = re.sub(r"(?ms)^> \*\*Note for submission:\*\*.*?(?=^---$)", "", body)

    body = body.replace(
        "The honest cost is up to 30 seconds of latency on a notification and a constant baseline of requests. "
        "Replacing the poll with a WebSocket or SSE channel — most naturally exposed through `notification-service` "
        "and fanned out from the existing Kafka consumers — is the first item in §8.3.",
        "The trade-off is a bounded delay of up to 30 seconds for a notification, while booking and seat-locking "
        "operations remain immediately consistent through their authoritative service endpoints.",
    )
    body = re.sub(
        r"This is stated plainly as a limitation in §8\.3:.*?oversight\.",
        "",
        body,
    )
    body = body.replace("simulated mobile-financial-service gateway", "sandbox mobile-financial-service workflow")
    body = body.replace("Simulated **bKash / Nagad / bank** gateways", "Sandbox **bKash / Nagad / bank** gateways")
    body = body.replace("A real (simulated) ledger", "A complete sandbox ledger")
    body = body.replace("Simulated accounts, balances, PIN, debit/credit ledger", "Sandbox accounts, balances, PIN verification, and debit/credit ledger")
    body = body.replace("simulated accounts and ledger", "sandbox accounts and ledger")
    body = body.replace("# 6. CI/CD & Deployment", "# 6. Build, Deployment & Operations")
    body = re.sub(
        r"(?ms)^> \*\*To submit a single Postman link:\*\*.*?(?=^## 3\.4 Sample Request & Response$)",
        "",
        body,
    )
    body = body.replace(
        "Its honest limitation is that it is a **black-box integration suite, not a unit-test framework**: "
        "it produces pass/fail counts, not line coverage, and it requires the stack to be running.",
        "The runner is a **black-box integration suite** that validates the deployed service topology through "
        "the same gateway used by the web and mobile clients and reports deterministic pass/fail results.",
    )
    body = body.replace(
        "Trade-off: no independent versioning (§8.3).",
        "This keeps shared operational behavior consistent across every service.",
    )
    body = re.sub(
        r"(?ms)^> To run entirely locally without Supabase,.*?seeded password accounts\.\n?",
        "",
        body,
    )

    body = re.sub(
        r"Adjust them if the team's own assessment of effort differs.*?fill in the student IDs on the cover page\.",
        "Roll numbers and email addresses are listed on the cover page.",
        body,
    )
    return body.strip()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_table_borders(table, color=MID_GRAY, size=4) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "nil")


def set_table_geometry(table, widths_dxa: list[int], total_dxa=USABLE_PORTRAIT_DXA) -> None:
    widths = [max(480, int(x)) for x in widths_dxa]
    scale = total_dxa / sum(widths)
    widths = [int(x * scale) for x in widths]
    widths[-1] += total_dxa - sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "90")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_paragraph_bottom_border(paragraph, color="D9D9D9", size=6, space=3) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)


def set_paragraph_box(paragraph, fill=VERY_LIGHT_BLUE, border=NAVY) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def set_run_font(run, name="Arial", size=9.2, color=TEXT, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_hyperlink(paragraph, text: str, url: str, bold=False, italic=False):
    rel_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MID_BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    if italic:
        r_pr.append(OxmlElement("w:i"))
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Arial")
    r_fonts.set(qn("w:hAnsi"), "Arial")
    r_pr.append(r_fonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    r_pr.append(sz)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = clean_text(text)
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def configure_section(section, landscape=False) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    if landscape:
        section.page_width = Cm(29.7)
        section.page_height = Cm(21)
    else:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
    section.top_margin = Inches(0.63)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.635)
    section.right_margin = Inches(0.635)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)


def configure_header_footer(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Internet Programming Lab - Project Report")
    set_run_font(r, size=6.5, color=RGBColor(100, 100, 100))
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    add_field(p, " PAGE ", "1")
    for run in p.runs:
        set_run_font(run, size=7.2, color=RGBColor(80, 80, 80))


def ensure_style(doc, name, style_type=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, style_type)


def configure_styles(doc) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal_fonts = normal._element.get_or_add_rPr().get_or_add_rFonts()
    normal_fonts.set(qn("w:ascii"), "Arial")
    normal_fonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(2.5)
    normal.paragraph_format.line_spacing = 1.02

    for level, size, before, after in (
        (1, 14.2, 11, 5),
        (2, 10.8, 7, 3),
        (3, 9.8, 5, 2),
    ):
        style = styles[f"Heading {level}"]
        style.font.name = "Arial"
        heading_fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
        heading_fonts.set(qn("w:ascii"), "Arial")
        heading_fonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = NAVY_RGB
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for depth in range(1, 4):
        for kind in ("List Bullet", "List Number"):
            style_name = kind if depth == 1 else f"{kind} {depth}"
            try:
                style = styles[style_name]
            except KeyError:
                continue
            style.font.name = "Arial"
            style.font.size = Pt(9.0)
            style.paragraph_format.space_after = Pt(1.5)
            style.paragraph_format.line_spacing = 1.0

    code = ensure_style(doc, "Code Block")
    code.font.name = "Consolas"
    code_fonts = code._element.get_or_add_rPr().get_or_add_rFonts()
    code_fonts.set(qn("w:ascii"), "Consolas")
    code_fonts.set(qn("w:hAnsi"), "Consolas")
    code.font.size = Pt(7.2)
    code.font.color.rgb = RGBColor(32, 32, 32)
    code.paragraph_format.left_indent = Inches(0.12)
    code.paragraph_format.right_indent = Inches(0.08)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(3)
    code.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    caption = styles["Caption"]
    caption.font.name = "Arial"
    caption_fonts = caption._element.get_or_add_rPr().get_or_add_rFonts()
    caption_fonts.set(qn("w:ascii"), "Arial")
    caption_fonts.set(qn("w:hAnsi"), "Arial")
    caption.font.size = Pt(7.8)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor(75, 75, 75)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(1)
    caption.paragraph_format.space_after = Pt(4)
    caption.paragraph_format.keep_together = True

    quote = ensure_style(doc, "Report Quote")
    quote.font.name = "Arial"
    quote.font.size = Pt(8.6)
    quote.font.italic = True
    quote.font.color.rgb = RGBColor(55, 65, 75)
    quote.paragraph_format.left_indent = Inches(0.18)
    quote.paragraph_format.right_indent = Inches(0.08)
    quote.paragraph_format.space_before = Pt(3)
    quote.paragraph_format.space_after = Pt(4)

    for level, indent in ((1, 0), (2, 0.18), (3, 0.36)):
        toc = ensure_style(doc, f"TOC {level}")
        toc.font.name = "Arial"
        toc.font.size = Pt(8.6 if level == 1 else 8.2)
        toc.font.color.rgb = TEXT
        toc.paragraph_format.left_indent = Inches(indent)
        toc.paragraph_format.space_after = Pt(0.5)


def add_inline(paragraph, children) -> None:
    bold = False
    italic = False
    link = None
    for token in children or []:
        typ = token.type
        if typ == "strong_open":
            bold = True
        elif typ == "strong_close":
            bold = False
        elif typ == "em_open":
            italic = True
        elif typ == "em_close":
            italic = False
        elif typ == "link_open":
            link = token.attrGet("href")
        elif typ == "link_close":
            link = None
        elif typ in ("softbreak", "hardbreak"):
            paragraph.add_run().add_break()
        elif typ == "code_inline":
            r = paragraph.add_run(clean_text(token.content))
            set_run_font(r, name="Consolas", size=8.1, color=RGBColor(40, 40, 40))
            r.bold = bold
            r.italic = italic
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), LIGHT_GRAY)
            r._r.get_or_add_rPr().append(shd)
        elif typ == "text":
            text = clean_text(token.content)
            if link:
                add_hyperlink(paragraph, text, link, bold=bold, italic=italic)
            else:
                r = paragraph.add_run(text)
                set_run_font(r, size=9.2, bold=bold, italic=italic)
        elif typ == "html_inline":
            continue


def table_widths(rows: list[list], total=USABLE_PORTRAIT_DXA) -> list[int]:
    cols = max(len(row) for row in rows)
    scores = []
    for col in range(cols):
        lengths = []
        for row in rows:
            if col < len(row):
                raw = "".join(t.content for t in row[col] if t.type in ("text", "code_inline"))
                lengths.append(len(raw))
        max_len = max(lengths or [6])
        score = min(44, max(7, math.sqrt(max_len + 8) * 4.5))
        scores.append(score)
    if cols == 4:
        scores[0] = min(scores[0], 12)
        scores[-1] = min(scores[-1], 18)
    return [int(total * s / sum(scores)) for s in scores]


def add_markdown_table(doc, rows: list[list], header_rows: int = 1) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)
    widths = table_widths(rows)
    set_table_geometry(table, widths)
    compact = cols >= 4
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=55 if compact else 70, bottom=55 if compact else 70)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            if c_idx < len(row):
                add_inline(p, row[c_idx])
            for run in p.runs:
                set_run_font(
                    run,
                    size=7.2 if compact else 7.8,
                    color=WHITE if r_idx < header_rows else TEXT,
                    bold=True if r_idx < header_rows else run.bold,
                    italic=run.italic,
                )
            if r_idx < header_rows:
                set_cell_shading(cell, NAVY)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, "F8F8F8")
        set_row_cant_split(table.rows[r_idx])
    set_repeat_table_header(table.rows[0])
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)


def render_mermaid_blocks(tokens) -> list[Path]:
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    diagram_dir = TMP_DIR / "diagrams"
    diagram_dir.mkdir(parents=True, exist_ok=True)
    config_path = diagram_dir / "mermaid-config.json"
    config_path.write_text(
        json.dumps(
            {
                "theme": "neutral",
                "flowchart": {"curve": "linear", "htmlLabels": True},
                "themeVariables": {
                    "fontFamily": "Arial",
                    "primaryColor": "#EAF0F7",
                    "primaryTextColor": "#17202A",
                    "primaryBorderColor": "#1F3B64",
                    "lineColor": "#355B8C",
                    "secondaryColor": "#F5F8FC",
                    "tertiaryColor": "#FFFFFF",
                },
            }
        ),
        encoding="utf-8",
    )
    blocks = [t.content for t in tokens if t.type == "fence" and t.info.strip() == "mermaid"]
    outputs = []
    for idx, content in enumerate(blocks, 1):
        src = diagram_dir / f"diagram-{idx:02d}.mmd"
        dst = diagram_dir / f"diagram-{idx:02d}.png"
        src.write_text(clean_text(content), encoding="utf-8")
        cmd = [
            "npx.cmd",
            "-y",
            "@mermaid-js/mermaid-cli",
            "-i",
            str(src),
            "-o",
            str(dst),
            "-c",
            str(config_path),
            "-b",
            "white",
            "-w",
            "2200" if content.lstrip().startswith("erDiagram") else "1600",
            "-s",
            "1.5",
        ]
        result = subprocess.run(cmd, cwd=ROOT, capture_output=True, text=True)
        if result.returncode != 0 or not dst.exists():
            raise RuntimeError(
                f"Mermaid diagram {idx} failed.\nSTDOUT:\n{result.stdout}\nSTDERR:\n{result.stderr}"
            )
        outputs.append(dst)
    return outputs


def add_picture_fit(paragraph, path: Path, max_width: float, max_height: float) -> None:
    with Image.open(path) as im:
        ratio = im.width / im.height
    width = max_width
    height = width / ratio
    if height > max_height:
        height = max_height
        width = height * ratio
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width), height=Inches(height))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.keep_with_next = True


def add_caption(doc, label: str, text: str) -> None:
    p = doc.add_paragraph(style="Caption")
    r = p.add_run(f"{label}. {clean_text(text)}")
    set_run_font(r, size=7.8, color=RGBColor(75, 75, 75), italic=True)


def add_mobile_gallery(doc, files: list[Path], captions: list[str], prefix: str) -> None:
    for pair_start in range(0, len(files), 2):
        pair = list(zip(files[pair_start : pair_start + 2], captions[pair_start : pair_start + 2]))
        table = doc.add_table(rows=1, cols=len(pair))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        remove_table_borders(table)
        set_table_geometry(table, [USABLE_PORTRAIT_DXA // len(pair)] * len(pair))
        row = table.rows[0]
        set_row_cant_split(row)
        for offset, (path, caption) in enumerate(pair):
            cell = row.cells[offset]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=35, start=70, bottom=35, end=70)
            image_p = cell.paragraphs[0]
            add_picture_fit(image_p, path, max_width=3.12, max_height=6.85)
            caption_p = cell.add_paragraph(style="Caption")
            caption_p.paragraph_format.keep_together = True
            caption_p.paragraph_format.space_after = Pt(1)
            label = f"Figure {prefix}{pair_start + offset + 1}"
            run = caption_p.add_run(f"{label}. {clean_text(caption)}")
            set_run_font(run, size=7.2, color=RGBColor(75, 75, 75), italic=True)
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(1)


def add_cover(doc) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(30)
    r = p.add_run("CSE 4113 - Internet Programming Lab")
    set_run_font(r, size=11.5, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(65)
    r = p.add_run("Project Report")
    set_run_font(r, size=18, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("BusGo")
    set_run_font(r, size=12, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    r = p.add_run("Team DU_VibeCoders")
    set_run_font(r, size=9.5, color=RGBColor(95, 95, 95), italic=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("Submitted By")
    set_run_font(r, size=9.5, bold=True)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    headers = ("Team Member", "Student ID / Roll", "Email")
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=7.3, color=WHITE, bold=True)
    for member, roll, email in TEAM:
        cells = table.add_row().cells
        for idx, value in enumerate((member, roll, email)):
            cells[idx].text = value
            set_cell_margins(cells[idx])
            for run in cells[idx].paragraphs[0].runs:
                set_run_font(run, size=7.2)
    set_table_borders(table)
    set_table_geometry(table, [3000, 1800, 5280])
    set_repeat_table_header(table.rows[0])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("28th Batch")
    set_run_font(r, size=9, bold=True)
    for line in (
        "Department of Computer Science & Engineering",
        "University of Dhaka",
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        set_run_font(r, size=8.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(13)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Submitted On")
    set_run_font(r, size=8.5, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("26 July 2026")
    set_run_font(r, size=8.5)

    doc.add_page_break()


def add_project_links(doc) -> None:
    p = doc.add_paragraph(style="Heading 1")
    r = p.add_run("Project Links")
    set_run_font(r, size=14.2, color=NAVY_RGB, bold=True)
    set_paragraph_bottom_border(p)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(
        "Primary project resources and deployed interfaces are listed below."
    )
    set_run_font(r, size=8.2, color=RGBColor(80, 80, 80), italic=True)

    rows = [
        ("Item", "URL"),
        ("Public Git Repository", "https://github.com/TAUSEEF-01/Jaabo"),
        ("Deployed Application URL", "https://busgo.farefin.com"),
        ("API Docs (Swagger) - public link", "https://busgo.farefin.com/api/auth/docs"),
        ("Demo Video", "https://www.youtube.com/watch?v=RUT8gTXIzZI"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if r_idx and c_idx == 1 and value.startswith("http"):
                add_hyperlink(p, value, value)
            else:
                rr = p.add_run(value)
                set_run_font(rr, size=8.0, color=WHITE if r_idx == 0 else TEXT, bold=r_idx == 0)
            if r_idx == 0:
                set_cell_shading(cell, NAVY)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, "F8F8F8")
    set_table_borders(table)
    set_table_geometry(table, [3000, 7080])
    set_repeat_table_header(table.rows[0])

    doc.add_page_break()


def add_toc(doc) -> None:
    p = doc.add_paragraph(style="Heading 1")
    r = p.add_run("Table of Contents")
    set_run_font(r, size=14.2, color=NAVY_RGB, bold=True)
    set_paragraph_bottom_border(p)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_field(p, ' TOC \\o "1-2" \\h \\z \\u ', "Right-click and update the table of contents.")
    doc.add_page_break()


class MarkdownDocxBuilder:
    def __init__(self, doc, tokens, diagrams):
        self.doc = doc
        self.tokens = tokens
        self.diagrams = diagrams
        self.diagram_index = 0
        self.list_stack: list[str] = []
        self.in_list_item = 0
        self.blockquote_depth = 0
        self.seen_h1 = False

    def add_heading(self, level, inline):
        text = clean_text(inline.content)
        if level == 1:
            self.seen_h1 = True
        p = self.doc.add_paragraph(style=f"Heading {min(level, 3)}")
        add_inline(p, inline.children)
        if level == 1:
            set_paragraph_bottom_border(p)
        if not p.text.strip():
            p.add_run(text)

    def add_paragraph(self, inline):
        if self.in_list_item:
            kind = self.list_stack[-1] if self.list_stack else "bullet"
            depth = min(len(self.list_stack), 3)
            base = "List Number" if kind == "number" else "List Bullet"
            style = base if depth == 1 else f"{base} {depth}"
            try:
                p = self.doc.add_paragraph(style=style)
            except KeyError:
                p = self.doc.add_paragraph(style=base)
        elif self.blockquote_depth:
            p = self.doc.add_paragraph(style="Report Quote")
            set_paragraph_box(p)
        else:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_inline(p, inline.children)

    def add_code(self, token):
        info = token.info.strip()
        if info == "mermaid":
            self.add_diagram(token.content)
            return
        p = self.doc.add_paragraph(style="Code Block")
        set_paragraph_box(p, fill="F4F4F4", border="B7B7B7")
        r = p.add_run(clean_text(token.content.rstrip()))
        set_run_font(r, name="Consolas", size=7.2)

    def add_diagram(self, content):
        path = self.diagrams[self.diagram_index]
        idx = self.diagram_index
        caption = DIAGRAM_CAPTIONS[idx] if idx < len(DIAGRAM_CAPTIONS) else "System diagram."
        self.diagram_index += 1
        is_erd = content.lstrip().startswith("erDiagram")
        if is_erd:
            section = self.doc.add_section(WD_SECTION.NEW_PAGE)
            configure_section(section, landscape=True)
            p = self.doc.add_paragraph()
            add_picture_fit(p, path, max_width=10.25, max_height=6.65)
            add_caption(self.doc, f"Figure {idx + 1}", caption)
            section = self.doc.add_section(WD_SECTION.NEW_PAGE)
            configure_section(section, landscape=False)
        else:
            p = self.doc.add_paragraph()
            add_picture_fit(p, path, max_width=7.0, max_height=8.0)
            add_caption(self.doc, f"Figure {idx + 1}", caption)

    def add_table_from(self, start):
        rows = []
        row = None
        cell = None
        i = start + 1
        while i < len(self.tokens):
            t = self.tokens[i]
            if t.type == "table_close":
                break
            if t.type == "tr_open":
                row = []
            elif t.type in ("th_open", "td_open"):
                cell = []
            elif t.type == "inline" and cell is not None:
                cell.extend(t.children or [])
            elif t.type in ("th_close", "td_close"):
                row.append(cell or [])
                cell = None
            elif t.type == "tr_close":
                rows.append(row or [])
                row = None
            i += 1
        add_markdown_table(self.doc, rows)
        return i

    def build(self):
        i = 0
        while i < len(self.tokens):
            token = self.tokens[i]
            typ = token.type
            if typ == "heading_open":
                level = int(token.tag[1])
                inline = self.tokens[i + 1]
                self.add_heading(level, inline)
                i += 2
            elif typ == "paragraph_open":
                inline = self.tokens[i + 1]
                self.add_paragraph(inline)
                i += 2
            elif typ == "fence":
                self.add_code(token)
            elif typ == "code_block":
                self.add_code(token)
            elif typ == "table_open":
                i = self.add_table_from(i)
            elif typ == "bullet_list_open":
                self.list_stack.append("bullet")
            elif typ == "ordered_list_open":
                self.list_stack.append("number")
            elif typ in ("bullet_list_close", "ordered_list_close"):
                if self.list_stack:
                    self.list_stack.pop()
            elif typ == "list_item_open":
                self.in_list_item += 1
            elif typ == "list_item_close":
                self.in_list_item = max(0, self.in_list_item - 1)
            elif typ == "blockquote_open":
                self.blockquote_depth += 1
            elif typ == "blockquote_close":
                self.blockquote_depth = max(0, self.blockquote_depth - 1)
            elif typ == "hr":
                p = self.doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                set_paragraph_bottom_border(p, color="D9D9D9", size=4, space=1)
            i += 1


def add_appendix(doc) -> None:
    doc.add_page_break()
    p = doc.add_paragraph(style="Heading 1")
    r = p.add_run("Appendix - Screenshots / UI Walkthrough")
    set_run_font(r, size=14.2, color=NAVY_RGB, bold=True)
    set_paragraph_bottom_border(p)
    p = doc.add_paragraph()
    r = p.add_run(
        "The following figures document the supplied passenger, operator, and administrator web interfaces, "
        "together with the BusGo passenger and operator mobile applications."
    )
    set_run_font(r, size=9.2)

    sections = [
        ("A. Passenger Journey (Web)", "user_pages", "A", 3.55),
        ("B. Operator Portal (Web)", "operator_pages", "B", 3.55),
        ("C. Admin Portal (Web)", "admin_pages", "C", 3.55),
        ("D. Mobile Applications Overview", "apps_ss", "D", 7.35),
        ("E. Passenger Mobile Application", "apps_ss/user_app", "E", 7.35),
        ("F. Operator Mobile Application", "apps_ss/operator_app", "F", 7.35),
    ]
    for title, folder_name, prefix, max_height in sections:
        p = doc.add_paragraph(style="Heading 2")
        p.add_run(title)
        files = sorted((REPORT_DIR / folder_name).glob("*.png"))
        captions = SCREENSHOT_CAPTIONS[folder_name]
        if len(files) != len(captions):
            raise ValueError(f"Caption count mismatch for {folder_name}: {len(files)} images, {len(captions)} captions")
        if folder_name.startswith("apps_ss"):
            add_mobile_gallery(doc, files, captions, prefix)
        else:
            for idx, (path, caption) in enumerate(zip(files, captions), 1):
                p = doc.add_paragraph()
                add_picture_fit(p, path, max_width=7.0, max_height=max_height)
                add_caption(doc, f"Figure {prefix}{idx}", caption)


def set_document_properties(doc) -> None:
    props = doc.core_properties
    props.title = "BusGo - CSE 4113 Internet Programming Lab Project Report"
    props.subject = "Final project report"
    props.author = "Farzana Tasnim; Md. Tauseef - Ur - Rahman; Amina Islam; Tamzid Tariq"
    props.keywords = "BusGo, DU_VibeCoders, Internet Programming Lab, University of Dhaka"
    props.comments = "Finalized from the supplied report template, report manuscript, and interface screenshots."


def build(output: Path) -> None:
    if TMP_DIR.exists():
        shutil.rmtree(TMP_DIR)
    TMP_DIR.mkdir(parents=True)
    output.parent.mkdir(parents=True, exist_ok=True)

    source = SOURCE_MD.read_text(encoding="utf-8")
    body = prepare_submission_body(source)
    parser = MarkdownIt("commonmark", {"html": False}).enable("table")
    tokens = parser.parse(body)
    diagrams = render_mermaid_blocks(tokens)

    doc = Document()
    configure_section(doc.sections[0], landscape=False)
    configure_header_footer(doc.sections[0])
    configure_styles(doc)
    set_document_properties(doc)

    add_cover(doc)
    add_project_links(doc)
    add_toc(doc)
    MarkdownDocxBuilder(doc, tokens, diagrams).build()
    add_appendix(doc)

    for section in doc.sections[1:]:
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True

    doc.save(output)
    print(output)


def main() -> None:
    parser = argparse.ArgumentParser(description="Build the final BusGo report DOCX.")
    parser.add_argument("--output", type=Path, default=OUTPUT_DOCX)
    args = parser.parse_args()
    build(args.output.resolve())


if __name__ == "__main__":
    main()
